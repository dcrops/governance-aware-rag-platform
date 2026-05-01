"""
Document ingestion supporting TXT, PDF, and DOCX.
Handles file validation, extraction, and returns a Document model.
"""

import hashlib
import os

from app.models.document import Document


def _extract_text_from_txt(file_path: str) -> str:
    """Extract raw text from a UTF-8 plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        raise IOError(f"Failed to read TXT file as UTF-8: {e}")


def _extract_text_from_pdf(file_path: str) -> tuple[str, list[dict]]:
    """
    Extract raw text content from a PDF file using pypdf.
    Returns:
        raw_text (str): Concatenated text from all non-empty pages.
        pages (list[dict]): List of {"page_number": int, "text": str} dicts for each non-empty page.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError(
            "pypdf library is required for PDF ingestion. "
            "Install with `pip install pypdf`."
        )

    try:
        reader = PdfReader(file_path)
        pages = []
        texts = []

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    page_dict = {
                        "page_number": i + 1,
                        "text": page_text
                    }
                    pages.append(page_dict)
                    texts.append(page_text)
            except Exception as e:
                raise RuntimeError(
                    f"Failed to extract text from PDF page {i + 1}: {e}"
                )
        raw_text = "\n\n".join(texts)
        return raw_text, pages
    except Exception as e:
        raise RuntimeError(f"Failed to process PDF file: {e}")



def _extract_text_from_docx(file_path: str) -> str:
    """Extract raw text content from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "python-docx library is required for DOCX ingestion. "
            "Install with `pip install python-docx`."
        )

    try:
        doc = docx.Document(file_path)

        texts = [p.text for p in doc.paragraphs if p.text]

        return "\n".join(texts)

    except Exception as e:
        raise IOError(f"Failed to extract text from DOCX: {e}")


def ingest_document(file_path: str) -> Document:
    """
    Ingest a document from disk.
    Supports TXT, PDF, and DOCX.

    Returns a Document instance with file type, metadata, source path, and raw text.
    For PDFs, also includes per-page text in metadata["pages"].
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document path does not exist: {file_path}")

    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)

    if file_size == 0:
        raise ValueError("File is empty.")

    _, ext = os.path.splitext(file_name)
    ext = ext.lower().strip(".")

    metadata = {
        "file_name": file_name,
        "file_size": file_size,
        "file_type": ext,
    }

    if ext == "txt":
        raw_text = _extract_text_from_txt(file_path)

    elif ext == "pdf":
        raw_text, pages = _extract_text_from_pdf(file_path)
        metadata["pages"] = pages

    elif ext == "docx":
        raw_text = _extract_text_from_docx(file_path)

    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    if not raw_text or not raw_text.strip():
        raise ValueError("Failed to extract non-empty text from the document.")

    doc_id = hashlib.sha256(raw_text.encode("utf-8")).hexdigest()

    return Document(
        doc_id=doc_id,
        source_path=file_path,
        file_type=ext,
        metadata=metadata,
        raw_text=raw_text,
    )